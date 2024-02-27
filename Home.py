import matplotlib as mpl
import matplotlib.pyplot as plt
import os
from matplotlib import font_manager as fm
from matplotlib.font_manager import FontProperties
import io
from docx import Document
from io import BytesIO
from docx.shared import Pt , Cm
#from docx.oxml.ns import nsdecls
#from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import altair as alt
import numpy as np
import pandas as pd
import streamlit as st
#from streamlit_option_menu import option_menu   

st.set_option('deprecation.showPyplotGlobalUse', False)
script_dir = os.path.dirname(os.path.abspath(__file__))
thai_font_path = os.path.join("THSARABUN BOLD.TTF")
thai_font_prop = fm.FontProperties(fname=thai_font_path)

digit = int(2)
list_pie_chart = {}
list_boxplot={}
list_bar_chart_comma={}
list_bar_chart={}
list_stack_str={}
list_stack_num={}
#list_stack_bar={}
list_str_stack = {}
list_num_stack = {}
list_comment={}
list_time={}
list_non = {}

color_set1 = ['#FF5733', '#FFBD33', '#33FF57', '#33FFBD', '#5733FF', '#BD33FF', '#FF33BD', '#FF3364', '#FF3333', '#33FFBD']
color_set2 = ['#FF9966', '#66FF99', '#9966FF', '#FFFF66', '#66FFFF', '#FF66FF', '#FF6666', '#66FF66', '#6666FF', '#FFCC66']
color_set3 = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']
color_set4 = ['#e7ba52', '#e7cb94', '#843c39', '#ad494a', '#d6616b', '#e7969c', '#7b4173', '#a55194', '#ce6dbd', '#de9ed6']

def upload(A):
 if A is not None:
  y = A.name.split(".")[1]
  if 'xlsx' in y:
   df = pd.read_excel(A)
  elif 'csv' in y:
   df = pd.read_csv(A)
  df.fillna('ไม่ระบุ',inplace=True)
  df.replace('-','ไม่ระบุ',inplace=True)
  df.replace(' ','ไม่ระบุ',inplace=True)
  #st.dataframe(df)  
  return df

def check_count(A):
 for i in A:
  if A[i] > 1 :
   return False
 return True

def num_check(A):
 for i in set(A):
  if type(i) is str and i != 'ไม่ระบุ':
   return False
  else:
   continue
 return True

def check_comma(A):
 for i in A:
   if (', ' in str(i)):
     return True
 return False

def split_comma(A):
 res = []
 for i in A:
  res = res + i.split(", ")
 return res

def Count(A,removenan=True):
 if removenan and 'ไม่ระบุ' in A:
  A = [n for n in A if n != 'ไม่ระบุ']
 list_A = list(set(A))
 counta = dict()
 for i in list_A:
  counta[i] = A.count(i)
 return counta

def count_list(A,removenan=True):
 if removenan and 'ไม่ระบุ'in A:
  A = [n for n in A if n != 'ไม่ระบุ']
 list_A = list(set(A))
 count_dict = dict()
 len_A = len(A)
 for c in list_A:
  per = (A.count(c)/len_A)*100
  count_dict[str(c)] = {"count":A.count(c),"percent":round(per,digit)}
 sorted_dict = dict(sorted(count_dict.items(),reverse=True))
 return sorted_dict

def stat(A,removenan=True):
 if removenan and 'ไม่ระบุ' in A:
  A = [n for n in A if n != 'ไม่ระบุ']
 A = [5 if x == 'มากที่สุด'else(4 if x == 'มาก'else(3 if x == 'ปานกลาง' else(2 if x == 'น้อย' else(1 if x == 'น้อยที่สุด' else x)))) for x in A]
 mean = np.mean(A)
 sd = np.std(A)
 mean_sd = {'ค่าเฉลี่ย':round(mean,digit),'ส่วนเบี่ยงเบนมาตรฐาน':round(sd,digit)}
 return mean_sd

def change_num_to_text(A):
 x = []
 dict_change_num_to_text = {5: 'มากที่สุด', 4: 'มาก', 3: 'ปานกลาง', 2: "น้อย", 1: "ควรปรับปรุง", 'ไม่ระบุ': 'ไม่ระบุ'}
 for i in upload_df[A].values.tolist():
  x.append(dict_change_num_to_text[i])
 x.sort()
 return x

def create_table(data,doc):
 headers = data[0]
 rows = data[1:]
 df = pd.DataFrame(rows, columns=headers)
 
 table = doc.add_table(df.shape[0]+1, df.shape[1])
 for j in range(df.shape[-1]):
  table.cell(0, j).text = df.columns[j]
 for i in range(df.shape[0]):
  for k in range(df.shape[-1]):
   table.cell(i+1, k).text = str(df.values[i,k])
   
 return table

def table_comment(data, doc):
 headers = data[0]
 rows = data[1:]
 df = pd.DataFrame(rows, columns=headers)

 table = doc.add_table(df.shape[0] + 1, df.shape[1])
 for j in range(df.shape[-1]):
  table.cell(0, j).text = df.columns[j]
 for i in range(df.shape[0]):
  for k in range(df.shape[-1]):
            
   cell = table.cell(i + 1, k)
   if k == 0:  
    cell.text = f"- {df.values[i, k]}"
   else:
    cell.text = str(df.values[i, k])

def create_word_doc(Pie_chart,Box_chart,Com_bar,Bar_chart,St_str,St_num,Str_st,Num_st,
                    table_pie,table_box,table_comma,table_bar,table_str,table_num,str_table,num_table,upload_file,comment):
 doc = Document()

 heading = doc.add_heading(level=0)
 heading_run = heading.add_run('สรุปผล' + upload_file.name.split('.')[0])
 heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
 heading_run.bold = True
 heading_run.font.size = Pt(20)                       
 #--------------------------------------------------ตาราง
 for t_p in table_pie:
  df = create_table(t_p,doc)  
  df.style = 'Table Grid'                       
  doc.add_paragraph('\t')# ตารางวงกลม
     
 for t_box in table_box:
  df = create_table(t_box,doc)
  df.style = 'Table Grid' # ตารางกล่อง
  doc.add_paragraph('\t')
     
 for t_com in table_comma:
  df = create_table(t_com,doc)
  df.style = 'Table Grid'
  doc.add_paragraph('\t')# ตารางแท่ง
        
 for t_bar in table_bar:
  df = create_table(t_bar,doc)
  df.style = 'Table Grid'
  doc.add_paragraph('\t')# ตารางแท่ง
 
 for t_str in table_str:
  df = create_table(t_str,doc)
  df.style = 'Table Grid'
  doc.add_paragraph('\t')# ตารางแท่งต่อกับมีหัวใหญ่
     
 for t_num in table_num:
  df = create_table(t_num,doc)
  df.style = 'Table Grid' 
  doc.add_paragraph('\t')# ตารางแท่งต่อกับมีหัวใหญ่
     
 for str_t in str_table:
  df = create_table(str_t,doc)
  df.style = 'Table Grid'# ตารางแท่งต่อกับไม่มีหัวใหญ่
  doc.add_paragraph('\t')
     
 for num_t in num_table:
  df = create_table(num_t,doc)
  df.style = 'Table Grid'# ตารางแท่งต่อกับไม่มีหัวใหญ่ 
  doc.add_paragraph('\t')
     
 for ment in comment:   
  df = table_comment(ment,doc)
  doc.add_paragraph('\t')   
 #---------------------------------------------------------------------------ภาพ
 head_pic = doc.add_heading(level=0)
 head = head_pic.add_run('ภาพแผนภูมิ')
 head.bold = True
 head.font.size = Pt(16)                                              

 for pie in Pie_chart:
  #doc.add_picture(pie, height=Cm(10.16))#, width=Cm(15.24), height=Cm(10.16)
  paragraph_pie = doc.add_paragraph()
  run_pie = paragraph_pie.add_run()
  picture_pie = run_pie.add_picture(pie, height=Cm(10.16))
  paragraph_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
 for box in Box_chart:
  #doc.add_picture(box, height=Cm(10.16))
   paragraph_box = doc.add_paragraph()
   run_box = paragraph_box.add_run()
   picture_box = run_box.add_picture(box, height=Cm(10.16))
   paragraph_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
     
 for com in Com_bar:
  #doc.add_picture(com, height=Cm(10.16))
   paragraph_com = doc.add_paragraph()
   run_com = paragraph_com.add_run()
   picture_com = run_com.add_picture(com, height=Cm(10.16))
   paragraph_com.alignment = WD_ALIGN_PARAGRAPH.CENTER
     
 for bar in Bar_chart:
  #doc.add_picture(bar, height=Cm(10.16))
  paragraph_bar = doc.add_paragraph()
  run_bar = paragraph_bar.add_run()
  picture_bar = run_bar.add_picture(bar, height=Cm(10.16))
  paragraph_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER   
    
 for str in St_str:
  #doc.add_picture(str, width=Cm(15.24))
  paragraph_str = doc.add_paragraph()
  run_str = paragraph_str.add_run()
  picture_str = run_str.add_picture(str, width=Cm(15.24))
  paragraph_str.alignment = WD_ALIGN_PARAGRAPH.CENTER
     
 for num in St_num:
  #doc.add_picture(num, width=Cm(15.24))
  paragraph_num = doc.add_paragraph()
  run_num = paragraph_num.add_run()
  picture_num = run_num.add_picture(num, width=Cm(15.24))
  paragraph_num.alignment = WD_ALIGN_PARAGRAPH.CENTER   
     
 for Str in Str_st:
  #doc.add_picture(Str, width=Cm(15.24))
  paragraph_Str = doc.add_paragraph()
  run_Str = paragraph_Str.add_run()
  picture_Str = run_Str.add_picture(Str, width=Cm(15.24))
  paragraph_Str.alignment = WD_ALIGN_PARAGRAPH.CENTER
     
 for Num in Num_st:
  #doc.add_picture(Num, width=Cm(15.24))
  paragraph_Num = doc.add_paragraph()
  run_Num = paragraph_Num.add_run()
  picture_Num = run_Num.add_picture(Num, width=Cm(15.24))
  paragraph_Num.alignment = WD_ALIGN_PARAGRAPH.CENTER 
     
 doc.save('report.docx')
 return 'report.docx'

def pie_chart(data, key):
 labels = [key for key in data]
 counts = [data[key]['percent'] for key in data]
 fig,ax = plt.subplots()
 ax.pie(counts, labels=labels, autopct=f'%.{digit}f', textprops={'fontproperties': thai_font_prop, 'fontsize': 14})
 plt.title(key, fontproperties=thai_font_prop, fontsize=16)
 if '/' in key:
  key.split('/')[0]
 chart_pie = f"{key}.png"
 plt.savefig(chart_pie, bbox_inches='tight',dpi=300)#, bbox_inches='tight'
 st.pyplot()
 return chart_pie

def boxplot(data,key,removenan=True):
 if removenan and 'ไม่ระบุ' in data:
  data = [n for n in data if n != 'ไม่ระบุ']
 fig,ax = plt.subplots()
 plt.boxplot(data,showmeans=True)
 q1 = np.percentile(data,25)
 q3 = np.percentile(data,75)
 median = np.median(data)
 average = np.mean(data)
 outlier_above = q3+(q3-q1)*1.5
 outlier_below = q1-(q3-q1)*1.5
 c = [x for x in data if (x < outlier_below or x > outlier_above)]
 c = list(set(c))
 c.sort()
 if len(c) > 0:
  for i in range(len(c)):
   plt.text(1.1,c[i],f"Outlier_{i+1}:{c[i]:.{digit}f}")
   data = [x for x in data if x not in c]
   max = np.max(data)
   min = np.min(data)
   plt.text(1.1,max,f'Max:{max:.{digit}f}')
   plt.text(1.1,min,f'Min:{min:.{digit}f}')
 else:
  max = np.max(data)
  min = np.min(data)
  plt.text(1.1,max,f'Max:{max:.{digit}f}')
  plt.text(1.1,min,f'Min:{min:.{digit}f}')
    
 plt.text(0.7,q1,f'Q1: {q1:.{digit}f}')
 plt.text(0.7,q3,f'Q3: {q3:.{digit}f}')
 plt.text(1.1, median, f'Q2: {median:.{digit}f}')
 plt.text(0.7,average, f'Average: {average:.{digit}f}')
 plt.title(key,fontproperties=thai_font_prop, fontsize = 16)
 if '/' in key:
  key.split('/')[0]
 chart_box = f"{key}.png"
 plt.savefig(chart_box, bbox_inches='tight',dpi=300)
 st.pyplot()
 return chart_box

def bar_list_count(data,orther_number=1):
 values = [data[key] for key in data if (data[key] > orther_number ) and (key != "ไม่ระบุ")]
 values_orther =  [data[key] for key in data if (data[key] <= orther_number ) and (key != "ไม่ระบุ")]
 labels = [key for key in data if (data[key] > orther_number ) and (key != "ไม่ระบุ")]
 if len(values_orther)>0:
  values.append(sum(values_orther))
  labels.append('อื่น ๆ')
 if 'ไม่ระบุ' in data:
  values.append(data['ไม่ระบุ'])
  labels.append('ไม่ระบุ')
 return [labels, values]

def bar_chart_new(data,key,legend):
 values = data[1]
 fig,ax = plt.subplots()
 if legend == True:
  labels = range(1,len(data[0])+1)
  ax.set_xticks(labels)
  ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
  color=plt.rcParams['axes.prop_cycle'].by_key()['color']
  for i in range(len(data[0])):
   Legend = f'{i + 1}:{data[0][i]}'
   ax.bar(labels, values, label=Legend,color=color)#,color=color  
  ax.legend(bbox_to_anchor=(1, 0, 0.18, 1),prop=thai_font_prop,handlelength=0, fontsize=14) 
 else:
  Label = data[0]
  color=plt.rcParams['axes.prop_cycle'].by_key()['color']
  ax.set_xticklabels(Label, fontproperties=thai_font_prop, fontsize=14)
  ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
  ax.bar(Label,values,color=color)
 plt.title(key,fontproperties=thai_font_prop, fontsize=16)
 i = 0
 chart_bar = f"{i+1}.png"
 plt.savefig(chart_bar, bbox_inches='tight',dpi=300)
 st.pyplot()
 return chart_bar
 
def stacked_bar(data,key):
 fig,ax = plt.subplots()
 name = data.keys()
 data1 = data.values()

 ax.set_yticklabels(name, fontproperties=thai_font_prop, fontsize=16)
 d_f = pd.DataFrame(data1,index=name)
 
 d_f.plot.barh(stacked=True, figsize=(8,3), ax=ax).legend(bbox_to_anchor=(1, 0, 0.19, 1), prop=thai_font_prop, fontsize=16)
 plt.title(key,fontproperties=thai_font_prop, fontsize=16)
 if '/' in key:
  key.split('/')[0]
 chart_stack = f"{key}.png"
 plt.savefig(chart_stack, bbox_inches='tight',dpi=300)
 st.pyplot()
 return chart_stack

with st.sidebar:
 upload_file = st.sidebar.file_uploader(" ",type=["csv", "xlsx"]) 
 upload_df = upload(upload_file) 
st.image("header.png")
with st.expander('### :red[คำแนะนำ]',expanded=(upload_file is None)):#(upload_file is None)
 st.markdown('**:blue[1. โหลดไฟล์ออกจากแบบฟอร์มที่ได้ทำการสร้างแบบสอบถาม]**')
 st.image("download.png")
 st.markdown("""---""")
 st.markdown('**:blue[2. กด Browse files จากนั้นเลือกไฟล์แบบสอบถามที่ต้องการ]**')
 st.image("browse.png")
 st.markdown("""---""")
 st.markdown('**:blue[3. เลือกประเภทของแผนภูมิที่ต้องการ]**')
 st.image("type.png")
 st.markdown("""---""")
 st.markdown('**:blue[4. ปรับแต่งรายละเอียดแผนภูมิ]**')
 st.image("customize.png")
 st.markdown("""---""")
 st.markdown('**:blue[5. เมื่อได้รูปแบบที่ต้องการแล้วกดปุ่มดาวน์โหลดเพื่อโหลดไฟล์ docs]**')
 
#-------------------------------------------------แยกหัวข้อ----------------------------------------------------#
 
if upload_file is not None:
 Color = st.radio('ปรับแต่งสีกราฟ', ['ชุดสีที่ 1', 'ชุดสีที่ 2', 'ชุดสีที่ 3', 'ชุดสีที่ 4'], horizontal=True, index=0)
 if Color == 'ชุดสีที่ 1':
  plt.rcParams['axes.prop_cycle'] = plt.cycler(color=color_set1)
 elif Color == 'ชุดสีที่ 2':
  plt.rcParams['axes.prop_cycle'] = plt.cycler(color=color_set2)
 elif Color == 'ชุดสีที่ 3':
  plt.rcParams['axes.prop_cycle'] = plt.cycler(color=color_set3)    
 else:
  plt.rcParams['axes.prop_cycle'] = plt.cycler(color=color_set4)   
      
 list_topic_stackbar=[]
 list_stackbar=[]
 list_question = [h for h in upload_df]
 if ('Times' or 'ประทับเวลา') in list_question[0]:
  list_question.pop(0)
      
 for key in list_question:
  column = upload_df[key].values.tolist()
  len_column = len(column)
  x = Count(column)
      
  if all(value == "ไม่ระบุ" for value in column):
   list_non[key] = True
   continue   
       
  if 'time' in key:
   list_time[key] = True
   continue

  if '[' in key:
   list_stackbar.append(key)
   topic = key.split('[')[0]
   if topic not in list_topic_stackbar:
    list_topic_stackbar.append(topic)
   continue
 
  if num_check(column) and set(column).issubset({1,2,3,4,5}):
   list_num_stack[key]={'removenan':True}
   continue

  if not num_check(column) and set(column).issubset({'มากที่สุด','มาก','ปานกลาง','น้อย','น้อยที่สุด'}):
   list_str_stack[key]={'removenan':True}
   continue
    
  if check_comma(column):
   list_bar_chart_comma[key] = {'removenan':True,'orther_number':1,'legend':True}
   continue
   
  if num_check(column):
   list_boxplot[key] = {'removenan':True}
   continue
   
  if len(set(column)) < 6:
   list_pie_chart[key]={'removenan':True}
  else:
   list_bar_chart[key] = {'removenan':True,'orther_number':1,'legend':True}
 
 set_topic = set(list_topic_stackbar)
 for i in set_topic:
  col = []
  for n in list_stackbar:
   if i in n:
     col.append(n)
  Column = upload_df[col].values.tolist()
  sum_Column = sum(Column,[])
  if num_check(sum_Column)and set(sum_Column).issubset({1,2,3,4,5,'ไม่ระบุ'}):
   for key in col:
    list_stack_num[key]={'removenan':True}
  else:
   for key in col:
    list_stack_str[key]={'removenan':True}
#--------------------------------------------------------------- ทำปุ่มแสดงเงื่อนไขของแต่ละหัวข้อ
#pie chart แสดงเพิ่มว่า ใส่ ไม่ระบุ หรือไม่

if upload_file is not None:
 list_pie_keys = list(list_pie_chart.keys())
 list_box_keys = list(list_boxplot.keys())
 list_bar_keys = list(list_bar_chart.keys())
 list_comma_keys = list(list_bar_chart_comma.keys())
 list_str_keys = list(list_stack_str.keys())
 list_num_keys = list(list_stack_num.keys())
 list_stackn_keys = list(list_num_stack.keys())
 list_stacks_keys = list(list_str_stack.keys())
  
 tab1, tab2 = st.sidebar.tabs(['ประเภทแผนภูมิ', 'ปรับแต่งรายระเอียดแผนภูมิ'])
 with tab1:
  x = 1000
  endtext =""
  topic_long = st.radio('แสดงหัวข้อแบบย่อ', ['ใช่', 'ไม่ใช่'], horizontal=True)
  if topic_long =='ใช่':
   x=32
   endtext = "ฯ"
  numberitem = 0
   
  for topic in list_pie_keys:
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + topic[:x]+endtext
   p = st.radio(head_bulet, ['แผนภูมิวงกลม (Pie chart)', 'แผนภูมิแท่ง (Bar chart)'])#, horizontal=True
   st.text("")
   if p == 'แผนภูมิแท่ง (Bar chart)':
    list_bar_chart[topic]={'removenan':True,'orther_number':1,'legend':True}
    del list_pie_chart[topic]
        
  if list_pie_keys != list():
   st.markdown("""---""")
       
  for topic in list_box_keys:
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + topic[:x]+endtext
   box = st.radio(head_bulet,['แผนภูมิกล่อง (Boxplot)','แผนภูมิแท่ง (Bar chart)'])
   st.text("")
   if box == 'แผนภูมิแท่ง (Bar chart)':
    list_bar_chart[topic]={'removenan':True,'orther_number':1,'legend':True}
    del list_boxplot[topic]
  if list_box_keys != list():
   st.markdown("""---""")
       
  for topic in list_comma_keys:
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + topic[:x]+endtext
   st.write(head_bulet)
   #comma = st.radio(head_bulet,['แผนภูมิแท่ง'])
   st.text("")
       
  for topic in list_bar_keys:
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + topic[:x]+endtext
   key = st.radio(head_bulet, ['แผนภูมิแท่ง (Bar chart)','แผนภูมิวงกลม (Pie chart)','ข้อเสนอแนะ (comment)'])
   st.text("")
   if key == 'แผนภูมิวงกลม (Pie chart)':
    list_pie_chart[topic]={'removenan':True}
    if 'orther_number' in list_bar_chart[topic]:
     del list_bar_chart[topic]
     continue
   if key == 'ข้อเสนอแนะ (comment)':
    list_comment[topic] = {'removenan':True}
    if 'orther_number' in list_bar_chart[topic]:
     del list_bar_chart[topic]
     continue
  if list_bar_keys != list():
   st.markdown("""---""")

  for topic in list_num_keys:
   sub_word = topic.split(' [')[1]
   sub_word = sub_word.strip().replace(']', '')
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + sub_word[:x]+endtext
   #st.write(head_bulet)
   num_val = st.radio(head_bulet,['แผนภูมิแท่งแบบต่อกัน (Stacked bar)','แผนภูมิกล่อง (Boxplot)'])
   if num_val == 'แผนภูมิกล่อง (Boxplot)':
    list_boxplot[topic] = {'removenan':True}
    del list_stack_num[topic]
   st.text("")
       
  for topic in list_stackn_keys:
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + topic[:x]+endtext
   #st.write(head_bulet)
   stack_num_val = st.radio(head_bulet,['แผนภูมิแท่งแบบต่อกัน (Stacked bar)','แผนภูมิกล่อง (Boxplot)'])
   if stack_num_val == 'แผนภูมิกล่อง (Boxplot)':
    list_boxplot[topic] = {'removenan':True}
    del list_num_stack[topic]   
   st.text("") 
      
  for topic in list_str_keys:
   sub_word = topic.split(' [')[1]
   sub_word = sub_word.strip().replace(']', '')
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + sub_word[:x]+endtext
   st.write(head_bulet)
   #str_val = st.radio(head_bulet,['แผนภูมิแท่งแบบต่อกัน'], horizontal=True)
   st.text("")

  for topic in list_stacks_keys:
   numberitem = numberitem+1
   strnumberitem = str(numberitem)+')'
   head_bulet = strnumberitem + topic[:x]+endtext
   st.write(head_bulet)
   #stack_str_val = st.radio(head_bulet,['แผนภูมิแท่งแบบต่อกัน'], horizontal=True)
   st.text("")   
     
 with tab2:
  Number = 0
  Type = st.radio('ประเภท',['วงกลม','กล่อง','แท่ง','แท่งต่อกัน'],horizontal=True)
  st.markdown("""---""")
  if Type == 'วงกลม':
   for topic_pie in list_pie_chart:
    Number = Number+1
    strnumberitem = str(Number)+')'
    head_bulet = strnumberitem + topic_pie[:x]+endtext
    Pie = st.radio(head_bulet, ['ลบไม่ระบุ','เพิ่มไม่ระบุ'], horizontal=True)
    list_pie_chart[topic_pie] = {'removenan': True if Pie == 'ลบไม่ระบุ' else False}
    continue
  if Type == 'กล่อง':
   for topic_box in list_boxplot:
    Number = Number+1
    strnumberitem = str(Number)+')'
    head_bulet = strnumberitem + topic_box[:x]+endtext 
    box = st.radio(head_bulet,['ไม่เพิ่ม','เพิ่ม'],horizontal=True)
    list_boxplot[topic_box]={'showmeans': True if box == 'ลบไม่ระบุ' else False}
    continue
  if Type == 'แท่ง':
   for topic in list_bar_chart_comma:
    Number = Number+1
    strnumberitem = str(Number)+')'
    head_bulet = strnumberitem + topic[:x]+endtext +'(หัวข้อใดที่ประสงค์เพิ่มข้อมูลของผู้ไม่ตอบแบบสอบถามในกราฟ)'
    A = upload_df[topic].values.tolist()
    a = split_comma(A)
    b = Count(a)
    bar = st.radio(head_bulet, ['ไม่เพิ่ม', 'เพิ่ม'], horizontal=True)
    bar_legend = st.radio(topic[:x]+endtext,['เพิ่มคำอธิบาย','ลบคำอธิบาย'], horizontal=True)
    y = st.slider(topic[:x]+endtext+(จำนวนความถี่ขั้นต่ำของแต่ละกราฟที่ประสงค์ให้ปรากฎแท่งในกราฟแต่ละหัวข้อ)', 0, max(b.values()), 1, 1) 
    list_bar_chart_comma[topic] = {'removenan': True if bar == 'ลบไม่ระบุ' else False, 'orther_number': y ,
                                    'legend': True if bar_legend == 'เพิ่มคำอธิบาย' else False}
   for topic_bar in list_bar_chart:
    Number = Number+1
    strnumberitem = str(Number)+')'
    head_bulet = strnumberitem + topic_bar[:x]+endtext + '(หัวข้อใดที่ประสงค์เพิ่มข้อมูลของผู้ไม่ตอบแบบสอบถามในกราฟ)'
    c = Count(upload_df[topic_bar].values.tolist())
    Bar = st.radio(head_bulet + \n + '1', ['ไม่เพิ่ม', 'เพิ่ม'], horizontal=True)
    Bar_legend = st.radio(topic_bar[:x]+endtext,['เพิ่มคำอธิบาย','ลบคำอธิบาย'], horizontal=True)
    y = st.slider(topic_bar[:x]+endtext+'(จำนวนความถี่ขั้นต่ำของแต่ละกราฟที่ประสงค์ให้ปรากฎแท่งในกราฟแต่ละหัวข้อ)', 0, max(c.values()), 1, 1)
    list_bar_chart[topic_bar] = {'removenan': True if Bar == 'ลบไม่ระบุ' else False, 'orther_number': y,
                                  'legend': True if Bar_legend == 'เพิ่มคำอธิบาย' else False}
    continue
  if Type == 'แท่งต่อกัน':
   for topic_stack in list_num_keys:
    Number = Number+1
    strnumberitem = str(Number)+')'
    head_bulet = strnumberitem + topic_stack[:x]+endtext
    num = st.radio(head_bulet,['ข้อมูลเชิงปริมาณ','ข้อมูลเชิงคุณภาพ'], horizontal=True)
    st.text("")
    if num == 'ข้อมูลเชิงคุณภาพ':
     list_stack_str[topic_stack]={'removenan':True}
     del list_stack_num[topic_stack]
   for topic_num in list_stackn_keys:
    Number = Number+1
    strnumberitem = str(Number)+')'
    head_bulet = strnumberitem + topic_num[:x]+endtext
    st_num = st.radio(head_bulet,['ข้อมูลเชิงปริมาณ','ข้อมูลเชิงคุณภาพ'], horizontal=True)
    st.text("")
    if st_num == 'ข้อมูลเชิงคุณภาพ':
     list_str_stack[topic_num]={'removenan':True}
     del list_num_stack[topic_num]
#-----------------------------------------------tab ภาพแผนภูมิ -------------------------------------------------------#
Pie_chart = []
Com_bar = []
Bar_chart = []
Box_chart = []
St_str = []
St_num = []
Num_st = []
Str_st = []

table_pie = []
table_box = []
table_comma = []
table_bar = []
table_str = []
table_num = []
str_table = []
num_table = []
comment = []

dict_str_stack = dict()
dict_num_stack = dict()
dict_stack_bar = dict()
dict_stack_str = dict()
top_name = None
if upload_file is not None:
 tab1, tab2 = st.tabs(['ภาพแผนภูมิ', 'ข้อมูลแบบตาราง'])
 with tab1:
  with st.expander('แผนภูมิวงกลม',expanded=True):
   for p in list_pie_chart:
    pie_chart_path = pie_chart(count_list(upload_df[p].values.tolist(),list_pie_chart[p]['removenan']),p)
    Pie_chart.append(pie_chart_path)
     
  with st.expander('แผนภาพกล่อง',expanded=True):
   for b in list_boxplot:
    box_charts = boxplot(upload_df[b].values.tolist(),b)
    Box_chart.append(box_charts)
     
  with st.expander('แผนภูมิแท่ง',expanded=True):
   for a in list_bar_chart_comma:
    A = upload_df[a].values.tolist()
    v = split_comma(A)
    count_v = Count(v,list_bar_chart_comma[a]['removenan'])
    data = bar_list_count(count_v,list_bar_chart_comma[a]['orther_number'])
    data_legend = list_bar_chart_comma[a]['legend']
    bar_comma = bar_chart_new(data,a,data_legend)
    Com_bar.append(bar_comma)
     
   for i in list_bar_chart:
    list_com = upload_df[i].values.tolist()
    a = Count(list_com,list_bar_chart[i]['removenan'])
    data = bar_list_count(a,list_bar_chart[i]['orther_number'])
    data_legend = list_bar_chart[i]['legend']
    bar_charts = bar_chart_new(data,i,data_legend)
    Bar_chart.append(bar_charts)
     
  with st.expander('แผนภูมิแท่งแบบต่อกัน',expanded=True):
   for i in list_stack_num:
    mat = upload_df[i].values.tolist()
    mean_sd = stat(mat)
    a = change_num_to_text(i)
    topic_word, sub_word = i.split(' [')[:2]
    topic_word = topic_word.strip()
    sub_word = sub_word.strip().replace(']', '')
    if topic_word != top_name:
     top_name = topic_word
    A_l = count_list(a)
    for k in A_l:
     A_l[k] = A_l[k]['percent']
    if topic_word not in dict_num_stack:
     dict_num_stack[topic_word] = dict()
    dict_num_stack[topic_word][sub_word] = A_l
   for i in dict_num_stack:
    stack_num = stacked_bar(dict_num_stack[i],i)
    St_num.append(stack_num)
     
   for i in list_num_stack:
    a = change_num_to_text(i)
    c = count_list(a)
    for k in c:
     c[k] = c[k]['percent']
    if i not in dict_stack_bar:
     dict_stack_bar[i] = dict()
    dict_stack_bar[i][''] = c
   for i in dict_stack_bar:
    num_st = stacked_bar(dict_stack_bar[i],i)
    Num_st.append(num_st)

   for i in list_stack_str:
    topic_word, sub_word = i.split(' [')[:2]
    topic_word = topic_word.strip()
    sub_word = sub_word.strip().replace(']','')
    A_l = count_list(upload_df[i].values.tolist(),list_stack_str[i]['removenan'])
    for k in A_l:
     A_l[k] = A_l[k]['percent']
    if topic_word not in dict_str_stack:
     dict_str_stack[topic_word] = dict()
    dict_str_stack[topic_word][sub_word] = A_l
   for s in dict_str_stack:
    stack_str = stacked_bar(dict_str_stack[s],s)
    St_str.append(stack_str)
     
   for i in list_str_stack:
    A_l = count_list(upload_df[i].values.tolist())
    for k in A_l:
     A_l[k] = A_l[k]['percent']
    if i not in dict_stack_str:
     dict_stack_str[i] = dict()
    dict_stack_str[i][''] = A_l
   for s in dict_stack_str:
    str_st = stacked_bar(dict_stack_str[s],s)
    Str_st.append(str_st)
     
  #----------------------------------------------------------------------------------------------------------------- tab2   
 with tab2:
  top_name = ''
  head_quality = ['หัวข้อ' , 'จำนวน' , 'เปอร์เซ็นต์']
  head_amount = ['หัวข้อ' , 'ค่าเฉลี่ย' , 'ส่วนเบี่ยงเบนมาตรฐาน']
  head_data = [['หัวข้อ', 'มากที่สุด','มาก','ปานกลาง','น้อย','น้อยที่สุด'],[" ", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)"]]
  head_re = ['หัวข้อ' , 'ค่าเฉลี่ย','ส่วนเบี่ยงเบนมาตรฐาน','แปรผล']
  data_pie = []
  data_box = []
  data_comma = []
  data_bar = []
  data_stack_str = []
  data_str_stack = []
  data_stack_num = []
  data_num_stack = []
  data_comment = []
  if list_pie_chart != dict() and {'removenan':True}: 
   st.markdown('<h3 style="color:blue; font-size:18px; text-align:center;">แผนภูมิวงกลม</h3>', unsafe_allow_html=True)
  for pie in list_pie_chart:
   values = count_list(upload_df[pie].values.tolist(), list_pie_chart[pie]['removenan'])
   data_pie.append([pie, 'จำนวน', 'เปอร์เซนต์'])
   for ans in values:
    count = values[ans]['count']
    percent = values[ans]['percent']
    data_pie.append([ans, count, percent,])
   data_pie.append(['รวม', sum([values[key]['count'] for key in values]), 100])
   table_pie.append(data_pie)
    

   #------------------------------ อ.เอกเขียนไว้
   st.markdown(f'<h3 style="color:red; font-size:16px">{pie}</h3>', unsafe_allow_html=True)
   st.table(data_pie)
   data_pie = []
   #---------------------------------------------

   
   #data_pie.append(['','','']) 
  if list_pie_chart != dict() and {'removenan':True}:
   st.markdown("""---""")

  if list_boxplot != dict() and {'removenan':True}:    
   st.markdown('<h3 style="color:blue; font-size:18px; text-align:center;">แผนภูมิกล่อง</h3>', unsafe_allow_html=True)  
  for box in list_boxplot:
   mean_sd = stat(upload_df[box].values.tolist())
   for m_s in mean_sd:    
    mean = mean_sd['ค่าเฉลี่ย']
    std = mean_sd['ส่วนเบี่ยงเบนมาตรฐาน']
    level = '' 
    if mean >= 4.2:
     level = 'มากที่สุด'
    elif mean >= 3.4:
     level = 'มาก'
    elif mean >= 2.6:
     level = 'ปานกลาง'
    elif mean >= 1.8:
     level = 'น้อย'
    elif mean < 1.8:
     level = 'น้อยที่สุด'
   data_box.append([box,mean,std,level])
  if list_boxplot != dict() and {'removenan':True}:
   table_box.append([head_re] + data_box)   
   st.markdown(f'<h3 style="color:red; font-size:18px">{box}</h3>', unsafe_allow_html=True)
   st.table([head_re,*data_box])
   st.markdown("""---""") 

  if list_bar_chart_comma != dict() and {'removenan':True,'orther_number':1}: 
   st.markdown('<h3 style="color:blue; font-size:18px; text-align:center;">แผนภูมิแท่ง</h3>', unsafe_allow_html=True)
       
  elif list_bar_chart != dict() and {'removenan':True,'orther_number':1}:
   st.markdown('<h3 style="color:blue; font-size:18px; text-align:center;">แผนภูมิแท่ง</h3>', unsafe_allow_html=True)
    
  for comma in list_bar_chart_comma:
   topic = upload_df[comma].values.tolist()
   all_number = len(topic)
   list_free = split_comma(topic)
   #set_list = list(set(list_free))
   val = Count(list_free,list_bar_chart_comma[comma]['removenan'])
   data = bar_list_count(val, list_bar_chart_comma[comma]['orther_number'])
   data_comma.append([comma, 'จำนวน', 'เปอร์เซนต์'])
   data_dict = dict(zip(data[0],data[1]))
   for key in data_dict:
    cou = data_dict[key]
    percent = 100*cou/len(list_free)
    data_comma.append([key,data_dict[key],round(percent,digit)])
   data_comma.append(['รวม',all_number,100])
   table_comma.append(data_comma)

   st.markdown(f'<h3 style="color:red; font-size:18px">{comma}</h3>', unsafe_allow_html=True)
   st.table(data_comma)
   data_comma = []
  
  for bar in list_bar_chart:
   data = upload_df[bar].values.tolist()
   Val = Count(data,list_bar_chart[bar]['removenan'])
   sum_val = sum([Val[key] for key in Val])
   other = bar_list_count(Val,list_bar_chart[bar]['orther_number'])
   data_bar.append([bar,'จำนวน', 'เปอร์เซนต์'])
   data_dict = dict(zip(other[0],other[1]))
   for key in data_dict:
    cou = data_dict[key]
    percent = 100*cou/sum_val
    data_bar.append([key,data_dict[key],round(percent,digit)])
   data_bar.append(['รวม',sum_val,100])
   table_bar.append(data_bar)

   st.markdown(f'<h3 style="color:red; font-size:18px">{bar}</h3>', unsafe_allow_html=True)
   st.table(data_bar)
   data_bar = []
    
  if list_bar_chart != dict() and {'removenan':True,'orther_number':1}:
   st.markdown("""---""")

  st.markdown('<h3 style="color:blue; font-size:18px; text-align:center;">แผนภูมิแท่งแบบต่อกัน</h3>', unsafe_allow_html=True) 
  for num in list_stack_num:
   mat = upload_df[num].values.tolist()
   mean_sd = stat(mat,True)
   topic_word, sub_word = num.split(' [')[:2]
   topic_word = topic_word.strip()
   sub_word = sub_word.strip().replace(']', '')
   if topic_word != top_name:
    data_stack_num.append([topic_word, 'ค่าเฉลี่ย','ส่วนเบี่ยงเบนมาตรฐาน','แปรผล'])
    top_name = topic_word 
   for key in mean_sd:
    mean = mean_sd['ค่าเฉลี่ย']
    s_d = mean_sd['ส่วนเบี่ยงเบนมาตรฐาน']
    level = '' 
    if mean >= 4.2:
     level = 'มากที่สุด'
    elif mean >= 3.4:
     level = 'มาก'
    elif mean >= 2.6:
     level = 'ปานกลาง'
    elif mean >= 1.8:
     level = 'น้อย'
    elif mean < 1.8:
     level = 'น้อยที่สุด'
   data_stack_num.append([sub_word,mean,s_d,level])
   table_num.append(data_stack_num)
  #st.table(data_stack_num)
  if list_stack_num != dict() and {'removenan':True}:
   st.table(data_stack_num)
    
  for nums in list_num_stack:
   math = upload_df[nums].values.tolist()
   mean_sd = stat(math,True)
   for key in mean_sd:
    mean = mean_sd['ค่าเฉลี่ย']
    s_d = mean_sd['ส่วนเบี่ยงเบนมาตรฐาน']
    level = '' 
    if mean >= 4.2:
     level = 'มากที่สุด'
    elif mean >= 3.4:
     level = 'มาก'
    elif mean >= 2.6:
     level = 'ปานกลาง'
    elif mean >= 1.8:
     level = 'น้อย'
    elif mean < 1.8:
     level = 'น้อยที่สุด'
   data_num_stack.append([nums,mean,s_d,level])

  if list_num_stack != dict() and {'removenan':True}:
   num_table.append([head_re] + data_num_stack)
   st.table([head_re,*data_num_stack])

  for Str in list_stack_str:
   Col = upload_df[Str].values.tolist()
   count_string = count_list(Col,list_stack_str[Str]['removenan'])
   topic_word, sub_word = Str.split(' [')[:2]
   topic_word = topic_word.strip()
   sub_word = sub_word.strip().replace(']', '')
   if topic_word != top_name:
    data_stack_str.append([topic_word, 'มากที่สุด','มาก','ปานกลาง','น้อย','น้อยที่สุด'])
    data_stack_str.append([" ", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)", "จำนวน(เปอร์เซนต์)"])
    top_name = topic_word
   if set(Col).issubset({'มากที่สุด','มาก','ปานกลาง','น้อย','น้อยที่สุด','ไม่ระบุ'}):
    data_stack_str.append([sub_word,f"{count_string['มากที่สุด']['count']}({count_string['มากที่สุด']['percent']}%)"if 'มากที่สุด' in count_string else "0(0%)",
                            f"{count_string['มาก']['count']}({count_string['มาก']['percent']}%)"if 'มาก' in count_string else "0(0%)",
                            f"{count_string['ปานกลาง']['count']}({count_string['ปานกลาง']['percent']}%)"if 'ปานกลาง' in count_string else "0(0%)",
                            f"{count_string['น้อย']['count']}({count_string['น้อย']['percent']}%)"if 'น้อย' in count_string else "0(0%)",
                            f"{count_string['น้อยที่สุด']['count']}({count_string['น้อยที่สุด']['percent']}%)"if 'น้อยที่สุด' in count_string else "0(0%)"])
   else:
    data_stack_str.append([sub_word,f"{count_string['5']['count']}({count_string['5']['percent']}%)"if '5' in count_string else "0(0%)",
                            f"{count_string['4']['count']}({count_string['4']['percent']}%)"if '4' in count_string else "0(0%)",
                            f"{count_string['3']['count']}({count_string['3']['percent']}%)"if '3' in count_string else "0(0%)",
                            f"{count_string['2']['count']}({count_string['2']['percent']}%)"if '2' in count_string else "0(0%)",
                            f"{count_string['1']['count']}({count_string['1']['percent']}%)"if '1' in count_string else "0(0%)"])
     
   table_str.append(data_stack_str)
 
  if list_stack_str != dict() and {'removenan':True}:
   st.table(data_stack_str)
    
  for strs in list_str_stack:
   Col = upload_df[strs].values.tolist()
   count_string = count_list(Col,list_str_stack[strs]['removenan'])
   if set(Col).issubset({'มากที่สุด','มาก','ปานกลาง','น้อย','น้อยที่สุด','ไม่ระบุ'}):
    data_stack_str.append([strs,f"{count_string['มากที่สุด']['count']}({count_string['มากที่สุด']['percent']}%)"if 'มากที่สุด' in count_string else "0(0%)",
                            f"{count_string['มาก']['count']}({count_string['มาก']['percent']}%)"if 'มาก' in count_string else "0(0%)",
                            f"{count_string['ปานกลาง']['count']}({count_string['ปานกลาง']['percent']}%)"if 'ปานกลาง' in count_string else "0(0%)",
                            f"{count_string['น้อย']['count']}({count_string['น้อย']['percent']}%)"if 'น้อย' in count_string else "0(0%)",
                            f"{count_string['น้อยที่สุด']['count']}({count_string['น้อยที่สุด']['percent']}%)"if 'น้อยที่สุด' in count_string else "0(0%)"])
   else:
    data_stack_str.append([strs,f"{count_string['5']['count']}({count_string['5']['percent']}%)"if '5' in count_string else "0(0%)",
                            f"{count_string['4']['count']}({count_string['4']['percent']}%)"if '4' in count_string else "0(0%)",
                            f"{count_string['3']['count']}({count_string['3']['percent']}%)"if '3' in count_string else "0(0%)",
                            f"{count_string['2']['count']}({count_string['2']['percent']}%)"if '2' in count_string else "0(0%)",
                            f"{count_string['1']['count']}({count_string['1']['percent']}%)"if '1' in count_string else "0(0%)"])
    
    
  if list_str_stack != dict() and {'removenan':True}:
   str_table.append([head_data] + data_stack_str)
   st.table(head_data+data_stack_str)
   st.markdown("""---""")
       
  if list_comment != dict()and {'removenan':True}: 
   st.markdown('<h3 style="color:blue; font-size:18px; text-align:center;">ข้อเสนอแนะ</h3>', unsafe_allow_html=True) 
  for ment in list_comment:
   Val = Count(upload_df[ment].values.tolist(),list_comment[ment]['removenan'])
   sorted_items = sorted(Val.items(), key=lambda x: x[1], reverse=True)
   data_comment.append([ment,'จำนวน'])
   for ans in sorted_items:
    count = ans[1]
    if count > 1:   
     data_comment.append([ans[0], f'({count})'])
    else:    
     data_comment.append([ans[0], ''])
   comment.append(data_comment)     
   st.table(data_comment)
   data_comment = []
    
    
#--------------------------------------------------------doc----------------------#

if upload_file is not None:
 word_file_path = create_word_doc(Pie_chart,Box_chart,Com_bar,Bar_chart,St_str,St_num,Str_st,Num_st,
                                  table_pie,table_box,table_comma,table_bar,table_str,table_num,str_table,num_table,upload_file,comment)
 st.download_button(label="ดาวน์โหลด",data=open(word_file_path, "rb").read(),file_name="report.docx",mime="application/docx")

#st.write(table_pie)
 #for t_p in table_pie:
  #df = create_table(t_p)
  #st.write(df)
#if st.button('Generate'):
 #chart_paths = []
 #pie_data = []
 #table_pie = []
 #doc = create_word_doc("This is the text content of the document.")
 #for p in list_pie_chart:
  #Values = count_list(upload_df[p].values.tolist(),list_pie_chart[p]['removenan'])
  #pie_chart_path = Pie_chart(Values,p)
  #pie_data.append([p,])
  #chart_paths.append(pie_chart_path)
 #word_file_path = create_word_doc(chart_paths)
 #if upload_file is not None:
  #st.download_button(label="Download Report",data=open(word_file_path, "rb").read(),file_name="report.docx",mime="application/docx")
 
 #for p in list_pie_chart:
  #pie_chart_path = pie_chart(count_list(upload_df[p].values.tolist(),list_pie_chart[p]['removenan']),p)
  #doc.add_picture(pie_chart_path,width=Inches(3.5)) 
  #pie_chart_path = pie_chart(pie_chart_data, pie_chart_key, pie_chart_digit,)
  #doc.add_picture(io.BytesIO(base64.b64decode(encoded_image)))  # เพิ่มรูปภาพ pie chart เข้าไปในเอกสาร
 # Save the document to a BytesIO object
 
 #doc_buffer = io.BytesIO()
 #doc.save(doc_buffer)
 #doc.seek(0)
 #doc_buffer.seek(1)
 #st.write(doc_buffer)
 
 # Provide download link for the generated document
 #st.download_button(label="Download Word Document",data=doc_buffer,file_name="output.docx",
     #mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",key="word-doc-download")
